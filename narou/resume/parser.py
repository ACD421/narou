from __future__ import annotations

import io
from pathlib import Path

import docx
import pdfplumber

from ..schema import Resume
from .sections import extract_contact, extract_sections, extract_skills


class ParseError(Exception):
    pass


def _parse_pdf_bytes(data: bytes) -> str:
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = []
        for page in pdf.pages:
            txt = page.extract_text() or ""
            pages.append(txt)
    return "\n\n".join(pages)


def _parse_docx_bytes(data: bytes) -> str:
    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.rstrip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style and not text.isupper():
            parts.append(text.upper())
        else:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_bytes(data: bytes, filename: str) -> Resume:
    name = filename.lower()
    if name.endswith(".pdf"):
        text = _parse_pdf_bytes(data)
    elif name.endswith(".docx"):
        text = _parse_docx_bytes(data)
    elif name.endswith(".txt"):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ParseError(f"unsupported file type: {filename}")

    text = text.strip()
    if not text:
        raise ParseError("no text extracted")

    sections = extract_sections(text)
    contact = extract_contact(text)
    skills = extract_skills(sections)
    return Resume(
        raw_text=text,
        sections=sections,
        contact=contact,
        skills=skills,
        source_filename=filename,
    )


def parse_resume(path: str | Path) -> Resume:
    p = Path(path)
    if not p.exists():
        raise ParseError(f"file not found: {p}")
    return parse_bytes(p.read_bytes(), p.name)
