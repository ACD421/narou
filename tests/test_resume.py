import io

import docx

from narou.resume import parse_bytes
from narou.resume.sections import extract_contact, extract_sections, extract_skills


SAMPLE = """JANE SMITH
jane@example.com   (555) 123-4567   https://linkedin.com/in/janesmith

SUMMARY
Senior software engineer with 8 years of experience building distributed systems in Python and Go.

EXPERIENCE
Senior Engineer, Acme Inc.   Jan 2020 - Present
Led team of 5 engineers building real-time ingestion pipeline on AWS.
Delivered 40% latency reduction across critical services.

Engineer, Beta LLC   Jun 2016 - Dec 2019
Built backend services in Python/Django serving 10M daily users.

EDUCATION
BS Computer Science, State University, 2016

SKILLS
Python, Go, AWS, Kubernetes, PostgreSQL, Docker, Terraform
"""


def test_extract_sections_finds_standard_headers():
    sections = extract_sections(SAMPLE)
    assert "summary" in sections
    assert "experience" in sections
    assert "education" in sections
    assert "skills" in sections
    assert "Python" in sections["skills"]


def test_extract_contact_finds_fields():
    contact = extract_contact(SAMPLE)
    assert contact["email"] == "jane@example.com"
    assert "555" in contact["phone"]
    assert "linkedin.com/in/janesmith" in contact["linkedin"]
    assert contact["name"] == "JANE SMITH"


def test_extract_skills_splits_on_commas():
    sections = extract_sections(SAMPLE)
    skills = extract_skills(sections)
    assert "Python" in skills
    assert "Kubernetes" in skills
    assert len(skills) >= 5


def test_parse_bytes_txt():
    r = parse_bytes(SAMPLE.encode("utf-8"), "sample.txt")
    assert r.contact.get("email") == "jane@example.com"
    assert "summary" in r.sections
    assert r.skills


def test_parse_bytes_docx():
    doc = docx.Document()
    doc.add_paragraph("JANE SMITH")
    doc.add_paragraph("jane@example.com")
    doc.add_paragraph("")
    doc.add_paragraph("SUMMARY")
    doc.add_paragraph("Senior engineer with 10 years building systems.")
    doc.add_paragraph("SKILLS")
    doc.add_paragraph("Python, AWS, Docker")
    buf = io.BytesIO()
    doc.save(buf)

    r = parse_bytes(buf.getvalue(), "sample.docx")
    assert r.contact.get("email") == "jane@example.com"
    assert "summary" in r.sections
    assert "Python" in r.skills
